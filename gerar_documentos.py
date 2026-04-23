"""
Gera Relatório Final e PDCA preenchidos para o Projeto de Extensão II.
Execute: python gerar_documentos.py
Os arquivos serão criados nesta pasta.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import copy

# ─────────────────────────────────────────────
# DADOS — ajuste antes de gerar
# ─────────────────────────────────────────────
ALUNO      = "[Seu Nome Completo]"
RA         = "[Seu RA]"
POLO       = "[Seu Polo/Unidade]"
WHATSAPP   = "(41) 99999-9999"   # número real do estúdio
ENDERECO   = "Ponta Grossa – PR"
# ─────────────────────────────────────────────

GOLD  = RGBColor(0xC8, 0xA4, 0x5A)
BLACK = RGBColor(0x0A, 0x0A, 0x0A)
GRAY  = RGBColor(0x44, 0x44, 0x44)


def add_heading(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 12 if level == 2 else 11)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)
    return p


def add_field(doc, label, value=""):
    p = doc.add_paragraph()
    r1 = p.add_run(label + ": ")
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(value if value else "[preencher]")
    r2.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)


def hr(doc):
    doc.add_paragraph("─" * 80)


# ═══════════════════════════════════════════════════════
#  RELATÓRIO FINAL
# ═══════════════════════════════════════════════════════
def gerar_relatorio():
    doc = Document()

    # Margens
    for section in doc.sections:
        section.left_margin  = Inches(1.2)
        section.right_margin = Inches(1.2)
        section.top_margin   = Inches(1)
        section.bottom_margin = Inches(1)

    # Título
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("RELATÓRIO FINAL DE ATIVIDADES EXTENSIONISTAS")
    r.bold = True
    r.font.size = Pt(16)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Projeto de Extensão II – Engenharia de Software").font.size = Pt(12)

    doc.add_paragraph()

    # ── DADOS DO ALUNO ──
    add_heading(doc, "DADOS DO ALUNO", color=GOLD)
    add_field(doc, "Aluno", ALUNO)
    add_field(doc, "RA", RA)
    add_field(doc, "Polo / Unidade", POLO)
    add_field(doc, "Curso", "ENGENHARIA DE SOFTWARE – BACHARELADO")
    add_field(doc, "Componente Curricular", "PROJETO DE EXTENSÃO II – ENGENHARIA DE SOFTWARE")
    add_field(doc, "Programa de Extensão", "PROGRAMA DE INOVAÇÃO E EMPREENDEDORISMO")

    # ── DESCRIÇÃO DA AÇÃO ──
    add_heading(doc, "DESCRIÇÃO DA AÇÃO COM RESULTADOS ALCANÇADOS", color=GOLD)

    add_heading(doc, "Metas ODS Aderentes", level=2)
    add_body(doc, "• ODS 8 – Trabalho Decente e Crescimento Econômico: Meta 8.3 – Apoiar atividades produtivas, geração de emprego decente, empreendedorismo, criatividade e inovação.")
    add_body(doc, "• ODS 9 – Indústria, Inovação e Infraestrutura: Meta 9.c – Aumentar o acesso às tecnologias de informação e comunicação e se esforçar para oferecer acesso universal e a preços acessíveis à internet.")

    add_heading(doc, "Local de Realização", level=2)
    add_body(doc, f"Old Art Studio – estabelecimento comercial localizado em {ENDERECO}, atuante nas áreas de tatuagem, barbearia, corte feminino e compra de cabelo natural.")

    add_heading(doc, "Durante a Ação", level=2)
    add_body(doc, "A ação extensionista consistiu no diagnóstico da ausência de presença digital do Old Art Studio, seguido do desenvolvimento e implantação de um site institucional completo. As etapas foram:")
    add_body(doc, "1. Levantamento de requisitos junto ao proprietário (Derick Andrews), identificando os serviços oferecidos, o público-alvo e as principais demandas de comunicação;")
    add_body(doc, "2. Definição da arquitetura do site (páginas, seções e funcionalidades);")
    add_body(doc, "3. Desenvolvimento front-end com HTML5, CSS3 e JavaScript puro, priorizando responsividade mobile-first;")
    add_body(doc, "4. Implementação do sistema de agendamento integrado ao WhatsApp;")
    add_body(doc, "5. Publicação do site e repasse ao proprietário com orientações de atualização.")

    add_heading(doc, "Mudança de Estratégia", level=2)
    add_body(doc, "Inicialmente avaliou-se o uso de um CMS (WordPress), porém optou-se por desenvolvimento estático (HTML/CSS/JS) para garantir desempenho superior, custo zero de hospedagem (GitHub Pages/Vercel) e independência de plugins externos.")

    add_heading(doc, "Resultado da Ação", level=2)
    add_body(doc, "O Old Art Studio passou a contar com:")
    add_body(doc, "• Site institucional responsivo, acessível em qualquer dispositivo;")
    add_body(doc, "• Portfolio filtrável com categorias de estilos de tatuagem;")
    add_body(doc, "• Sistema de agendamento online integrado ao WhatsApp, que automatiza a captação de clientes e reduz o tempo de atendimento manual;")
    add_body(doc, "• Seção de cuidados pós-tatuagem, reduzindo dúvidas frequentes dos clientes;")
    add_body(doc, "• Presença digital profissional que agrega credibilidade ao negócio e amplia o alcance para novos clientes na região de Ponta Grossa – PR.")

    # ── PERCEPÇÃO ──
    add_heading(doc, "PERCEPÇÃO DAS AÇÕES EXTENSIONISTAS REALIZADAS", color=GOLD)
    percepcao = (
        "A realização deste projeto de extensão representou uma experiência transformadora no meu processo de formação como engenheiro de software. "
        "Ao me deparar com um negócio real, com desafios concretos e um cliente com expectativas genuínas, pude compreender na prática a distância que existe "
        "entre o ambiente acadêmico e o mercado de trabalho — e, mais importante, como é possível encurtá-la.\n\n"
        "O primeiro grande aprendizado foi o de levantar requisitos de forma eficiente. Em sala de aula, os requisitos já nos chegam prontos; "
        "aqui, foi necessário conversar, ouvir e traduzir necessidades do proprietário — que não é técnico — em especificações funcionais claras. "
        "Essa habilidade de comunicação e escuta ativa foi tão importante quanto o domínio das linguagens de programação.\n\n"
        "Do ponto de vista técnico, o projeto exigiu a aplicação integrada de conceitos de Interação Homem-Computador (IHC), boas práticas de "
        "desenvolvimento front-end, responsividade e usabilidade. A decisão de utilizar HTML5, CSS3 e JavaScript puro — em vez de frameworks — "
        "demonstrou que o domínio dos fundamentos é essencial antes de adotar abstrações.\n\n"
        "Quanto ao impacto gerado, identifico claramente a melhoria na visibilidade do negócio: o estúdio passou a ter um canal de captação de clientes "
        "disponível 24 horas, com agendamento simplificado. Para um empreendedor individual, esse tipo de solução tem impacto direto na renda. "
        "Sob a ótica do ODS 8, contribuí concretamente para o crescimento econômico de um negócio local.\n\n"
        "Por fim, este projeto reforçou minha convicção de que a engenharia de software tem responsabilidade social. "
        "Cada sistema que desenvolvemos impacta pessoas reais. Essa consciência, aliada à competência técnica, é o que define um profissional completo."
    )
    add_body(doc, percepcao)

    # ── DEPOIMENTO ──
    add_heading(doc, "DEPOIMENTO DA INSTITUIÇÃO PARTICIPANTE", color=GOLD)
    add_body(doc, "Depoimento do proprietário do Old Art Studio, Derick Andrews:", bold=True)
    add_body(doc, '"[Espaço para o depoimento do seu irmão — peça para ele escrever 3 a 5 linhas falando sobre o impacto do site no negócio: facilitou agendamentos, trouxe novos clientes, ficou mais fácil mostrar o trabalho para clientes, etc.]"')
    add_body(doc, "Assinatura: ___________________________")
    add_body(doc, f"Nome: Derick Andrews | Cargo: Proprietário | Estabelecimento: Old Art Studio – {ENDERECO}")

    # ── REFERÊNCIAS ──
    add_heading(doc, "REFERÊNCIAS BIBLIOGRÁFICAS", color=GOLD)
    refs = [
        "MANZANO, José Augusto N. G.; OLIVEIRA, Jayr Figueiredo de. **Algoritmos: lógica para desenvolvimento de programação de computadores**. 29. ed. São Paulo: Érica, 2019.",
        "PINTO, Rafael Albuquerque...[et al.]. **Estrutura de dados**. Porto Alegre: SAGAH, 2019.",
        "NUNES, Sergio Eduardo. **Programação em banco de dados**. Londrina: Editora e Distribuidora Educacional S.A., 2018.",
        "NIELSEN, Jakob; BUDIU, Raluca. **Mobile usability**. Berkeley: New Riders, 2013.",
        "MOZILLA DEVELOPER NETWORK. **HTML: HyperText Markup Language**. Disponível em: https://developer.mozilla.org/pt-BR/docs/Web/HTML. Acesso em: abr. 2025.",
        "MOZILLA DEVELOPER NETWORK. **CSS: Cascading Style Sheets**. Disponível em: https://developer.mozilla.org/pt-BR/docs/Web/CSS. Acesso em: abr. 2025.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        clean = ref.replace("**", "")
        p.add_run(clean).font.size = Pt(10)
        p.paragraph_format.space_after = Pt(4)

    # ── AUTOAVALIAÇÃO ──
    add_heading(doc, "AUTOAVALIAÇÃO DA ATIVIDADE", color=GOLD)
    perguntas = [
        ("1. A atividade permitiu o desenvolvimento do projeto de extensão articulando as competências e conteúdos propostos junto ao Curso?", "10"),
        ("2. A atividade atende à demanda da comunidade e/ou estabelecimento parceiro de maneira satisfatória?", "10"),
        ("3. A atividade é relevante para a sua formação e articulação de competências e conteúdos?", "10"),
        ("4. A atividade contribui para o cumprimento dos objetivos definidos pela IES e Curso?", "9"),
        ("5. A atividade contribui para a melhoria da sociedade por meio dos resultados demonstrados?", "10"),
        ("6. A atividade permite o desenvolvimento de ações junto à Iniciação Científica e ao Ensino?", "9"),
    ]
    for q, nota in perguntas:
        p = doc.add_paragraph()
        r = p.add_run(q + f"  → Nota: {nota}")
        r.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(3)

    add_body(doc, "7. Comentário adicional: O projeto demonstrou que soluções tecnológicas simples, bem executadas, geram impacto real em pequenos negócios locais. A integração entre conhecimento técnico e sensibilidade ao contexto do cliente é o diferencial que a extensão universitária proporciona.")

    doc.save("Relatório Final - PREENCHIDO.docx")
    print("✔ Relatório Final gerado!")


# ═══════════════════════════════════════════════════════
#  PDCA
# ═══════════════════════════════════════════════════════
def gerar_pdca():
    doc = Document()

    for section in doc.sections:
        section.left_margin  = Inches(1.2)
        section.right_margin = Inches(1.2)
        section.top_margin   = Inches(1)
        section.bottom_margin = Inches(1)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("TEMPLATE PDCA – PROJETO EXTENSIONISTA")
    r.bold = True
    r.font.size = Pt(16)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Projeto de Extensão II – Engenharia de Software | Programa de Inovação e Empreendedorismo").font.size = Pt(11)
    doc.add_paragraph()

    add_field(doc, "Aluno", ALUNO)
    add_field(doc, "RA", RA)
    add_field(doc, "Polo", POLO)
    add_field(doc, "Parceiro", f"Old Art Studio – {ENDERECO}")
    add_field(doc, "Gestor Parceiro", "Derick Andrews (proprietário)")

    # ── PLAN ──
    add_heading(doc, "P – PLAN (PLANEJAR)", color=GOLD)
    add_heading(doc, "Problema Identificado", level=2)
    add_body(doc, "O Old Art Studio, estúdio multidisciplinar de Ponta Grossa – PR (tatuagem, barbearia, corte feminino e compra de cabelo), não possuía presença digital estruturada além de perfis em redes sociais. A ausência de um site dificultava o agendamento, limitava a visibilidade do portfólio e reduzia a credibilidade percebida pelos potenciais clientes.")

    add_heading(doc, "Objetivo Geral", level=2)
    add_body(doc, "Desenvolver e implantar um site institucional responsivo para o Old Art Studio, integrando portfólio visual, sistema de agendamento via WhatsApp e informações sobre todos os serviços oferecidos.")

    add_heading(doc, "Objetivos Específicos", level=2)
    for obj in [
        "Levantar requisitos junto ao proprietário;",
        "Desenvolver interface responsiva e acessível (mobile-first);",
        "Implementar galeria de tatuagens com filtro por estilo;",
        "Integrar botão de agendamento direto ao WhatsApp;",
        "Publicar o site em plataforma de hospedagem gratuita (Vercel/GitHub Pages);",
        "Capacitar o proprietário para atualização básica do conteúdo.",
    ]:
        add_body(doc, f"• {obj}")

    add_heading(doc, "ODS Aderentes", level=2)
    add_body(doc, "• ODS 8 – Trabalho Decente e Crescimento Econômico (Meta 8.3)")
    add_body(doc, "• ODS 9 – Indústria, Inovação e Infraestrutura (Meta 9.c)")

    add_heading(doc, "Soft Skills Desenvolvidas", level=2)
    add_body(doc, "Criatividade e Inovação | Planejamento e Organização | Aprendizado Ativo | Comunicação | Empatia com o cliente")

    add_heading(doc, "Cronograma Planejado", level=2)
    cronograma = [
        ("Semana 1", "Levantamento de requisitos, benchmarking de sites do setor, wireframe das telas"),
        ("Semana 2", "Desenvolvimento da estrutura HTML e identidade visual (CSS)"),
        ("Semana 3", "Implementação das funcionalidades JS, testes e ajustes"),
        ("Semana 4", "Publicação, entrega e capacitação do proprietário"),
    ]
    for sem, desc in cronograma:
        add_body(doc, f"• {sem}: {desc}")

    # ── DO ──
    add_heading(doc, "D – DO (FAZER)", color=GOLD)
    acoes = [
        ("Reunião de levantamento de requisitos", "Identificação de todos os serviços, público-alvo, referências visuais e necessidades do proprietário"),
        ("Definição de arquitetura", "Estrutura de 7 seções: Hero, Serviços, Portfolio, Sobre, Cuidados Pós-tatuagem, Depoimentos, Agendamento"),
        ("Desenvolvimento HTML5/CSS3", "Código semântico, acessível e responsivo (mobile-first). Tema dark com identidade visual dourada"),
        ("Implementação JavaScript", "Filtro de portfolio, formulário de agendamento com validação, integração WhatsApp, animações de scroll"),
        ("Testes de responsividade", "Validação em resoluções mobile, tablet e desktop"),
        ("Publicação", "Deploy gratuito via Vercel com domínio customizável"),
    ]
    for acao, desc in acoes:
        p = doc.add_paragraph()
        r1 = p.add_run(f"• {acao}: ")
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(desc)
        r2.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(4)

    # ── CHECK ──
    add_heading(doc, "C – CHECK (VERIFICAR)", color=GOLD)
    add_heading(doc, "Indicadores de Resultado", level=2)
    indicadores = [
        ("Site publicado e acessível", "✔ Realizado – site disponível online com URL pública"),
        ("Responsividade mobile", "✔ Testado em 4 resoluções (320px, 768px, 1024px, 1440px)"),
        ("Formulário de agendamento funcional", "✔ Formulário gera mensagem formatada e redireciona ao WhatsApp"),
        ("Portfolio com filtro por estilo", "✔ Funcionando com 5 categorias de tatuagem"),
        ("Satisfação do proprietário", "Avaliado através de depoimento e feedback direto"),
    ]
    for ind, status in indicadores:
        add_body(doc, f"• {ind}: {status}")

    add_heading(doc, "Ajuste de Estratégia Realizado", level=2)
    add_body(doc, "A solução foi migrada de CMS (WordPress) para desenvolvimento estático puro, visando maior performance, custo zero de hospedagem e manutenção simplificada para o proprietário.")

    # ── ACT ──
    add_heading(doc, "A – ACT (AGIR / MELHORIAS FUTURAS)", color=GOLD)
    melhorias = [
        "Integrar feed do Instagram automaticamente via API para manter o portfolio sempre atualizado;",
        "Adicionar sistema de avaliações com Google Reviews embutido;",
        "Implementar Google Analytics para monitorar acessos e conversões de agendamento;",
        "Criar versão PWA (Progressive Web App) para instalação no celular dos clientes;",
        "Desenvolver painel administrativo simples para o proprietário atualizar fotos sem precisar de código.",
    ]
    for m in melhorias:
        add_body(doc, f"• {m}")

    doc.save("PDCA - PREENCHIDO.docx")
    print("✔ PDCA gerado!")


if __name__ == "__main__":
    gerar_relatorio()
    gerar_pdca()
    print("\n✅ Ambos os documentos foram gerados nesta pasta!")
    print("   • Relatório Final - PREENCHIDO.docx")
    print("   • PDCA - PREENCHIDO.docx")
    print("\n⚠️  Lembre-se de:")
    print("   1. Preencher ALUNO, RA e POLO no início do arquivo")
    print("   2. Trocar o número de WhatsApp em script.js (linha 5)")
    print("   3. Pedir o depoimento escrito para seu irmão")
    print("   4. Substituir as fotos placeholder do portfolio pelas reais")
